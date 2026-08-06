#!/usr/bin/env python3
"""Ingest one file of every supported type, end to end, and confirm each becomes
retrievable chunks."""
import io
import sys
import time

sys.path.insert(0, "/home/yaswant7/projects/SpaceLabs/spacebot")
from sb import chunks, db, retrieval  # noqa: E402
from sb.media import pipeline as mpipe  # noqa: E402

db.init_db()
chunks.init()
fails = 0


def make_docx():
    import docx
    d = docx.Document()
    d.add_heading("Expense Policy", 1)
    d.add_paragraph("Employees may claim reasonable travel costs incurred on company "
                    "business. Claims must be submitted within 30 days of the expense.")
    d.add_heading("Limits", 2)
    d.add_paragraph("Hotel stays are capped at 180 GBP per night in London and 120 GBP "
                    "elsewhere in the UK. Meals are capped at 40 GBP per day.")
    t = d.add_table(rows=3, cols=2)
    for r, (a, b) in enumerate([("Category", "Daily cap"), ("Meals", "40 GBP"),
                                ("Taxi", "60 GBP")]):
        t.rows[r].cells[0].text, t.rows[r].cells[1].text = a, b
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def make_xlsx():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Oncall"
    ws.append(["Week", "Primary", "Secondary", "Escalation"])
    ws.append(["2026-W31", "sarah", "arjun", "meena"])
    ws.append(["2026-W32", "arjun", "priya", "sarah"])
    ws.append(["2026-W33", "priya", "meena", "arjun"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


CSV = (b"Vendor,Status,Tax ID,Owner\n"
       b"Acme Ltd,approved,GB123456,meena\n"
       b"Globex,pending,GB998877,meena\n"
       b"Initech,rejected,,priya\n")

HTML = (b"<html><head><style>.x{color:red}</style><script>var a=1</script></head><body>"
        b"<nav>Home | Docs | Support</nav>"
        b"<h1>VPN access</h1>"
        b"<p>Connect to the corporate VPN using the Tailscale client, which is installed "
        b"through the company app catalogue. Sign in with your SpaceLabs Google account "
        b"rather than a personal one, or the device will register against the wrong "
        b"tailnet and you will need IT to remove it before you can retry.</p>"
        b"<p>Request the <code>eng</code> tag from IT before your first connection. "
        b"Without it the client authenticates successfully but routes no traffic, which "
        b"looks exactly like a broken network and is the most common support ticket we "
        b"get from new joiners.</p>"
        b"<p>Once connected you should be able to reach the internal Argo CD dashboard "
        b"and the staging cluster. If either times out, disconnect and reconnect before "
        b"raising a ticket, since the route table occasionally needs a refresh.</p>"
        b"<footer>Copyright SpaceLabs</footer></body></html>")

# Its own keys, never the demo corpus's. This test used to ingest into POLICY.EXPENSES,
# OPS.ONCALL_ROTA, PROC.VENDOR_LIST and IT.VPN_ACCESS — real entries — which overwrote their
# category and owner with "Test"/"tester" and replaced their model-extracted subjects with
# whatever that run happened to produce.
#
# A test that mutates the data it runs against corrupts the thing it exists to protect, and
# this reached the product: a user asking about the VPN saw "Category: Test · Owner:
# @tester" in the provenance under their answer.
CASES = [
    ("DOCX", "expense-policy.docx", make_docx(),
     "TEST.FMT_DOCX", "what is the hotel cap in London?"),
    ("XLSX", "oncall-rota.xlsx", make_xlsx(),
     "TEST.FMT_XLSX", "who is on call in week 2026-W32?"),
    ("CSV", "vendors.csv", CSV,
     "TEST.FMT_CSV", "is Globex an approved vendor?"),
    ("HTML", "vpn.html", HTML,
     "TEST.FMT_HTML", "how do I get on the VPN?"),
]

for label, fname, data, key, question in CASES:
    job = mpipe.start_ingest(key, "", "Test", "tester",
                             [{"filename": fname, "mime": "", "bytes": data}])
    for _ in range(240):
        j = db.get_job(job)
        if j["status"] in ("drafted", "failed"):
            break
        time.sleep(1)
    j = db.get_job(job)
    if j["status"] != "drafted":
        print(f"  FAIL {label:5} ingest {j['status']}: {j.get('error')}")
        fails += 1
        continue

    db.set_workflow_status(key, "published")
    card = next((c for c in db.get_catalog() if c["wf_key"] == key), None)
    n = chunks.reindex_workflow(card["id"]) if card else 0
    chunks.embed_pending()
    retrieval.invalidate()

    res = retrieval.retrieve(question)
    hit = any(c["wf_key"] == key for c in res["chunks"])
    ok = hit and res["evidence"] >= 0.15
    fails += 0 if ok else 1
    print(f"  {'ok  ' if ok else 'FAIL'} {label:5} {n:3} chunks  evidence={res['evidence']:<6} "
          f"retrieved={hit}")
    if hit:
        top = next(c for c in res["chunks"] if c["wf_key"] == key)
        print(f"        -> {top['text'][:96]!r}")

# Clean up after itself. Left behind, these accumulate on every run and start competing
# with real documents in retrieval — the corpus grew from 83 chunks to 115 over a day of
# test runs, and every one of those extra chunks was a near-duplicate of a demo document.
for _case in CASES:
    db.delete_workflow(_case[3])
retrieval.invalidate()

print(f"\n{len(CASES)-fails}/{len(CASES)} formats ingested and retrievable")
print("index:", chunks.stats())
sys.exit(1 if fails else 0)
