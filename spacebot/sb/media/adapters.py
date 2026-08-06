"""Source adapters — turn one raw upload into a list of Segment descriptors.

A descriptor is a plain dict the pipeline persists/enriches:
  {modality, text, image_bytes?, anchor, meta}

Add a modality = add an adapter here. Nothing downstream changes, because everything
downstream consumes Segments, not files. That is the whole point.
"""
import io
import re

from .capabilities import CapabilityUnavailable, get_frame_extractor, get_transcriber


def guess_kind(mime: str, filename: str) -> str:
    m = (mime or "").lower()
    f = (filename or "").lower()
    if "pdf" in m or f.endswith(".pdf"):
        return "pdf"
    if f.endswith((".docx", ".docm")) or "wordprocessingml" in m:
        return "docx"
    if f.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in m:
        return "sheet"
    if f.endswith((".csv", ".tsv")) or m in ("text/csv", "text/tab-separated-values"):
        return "csv"
    if f.endswith((".html", ".htm")) or "text/html" in m:
        return "html"
    if m.startswith("image/") or f.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff")):
        return "image"
    if m.startswith("video/") or f.endswith((".mp4", ".mov", ".webm", ".mkv", ".avi")):
        return "video"
    if m.startswith("audio/") or f.endswith((".mp3", ".wav", ".m4a", ".ogg", ".flac")):
        return "audio"
    return "transcript"


class Adapter:
    def decompose(self, data: bytes, filename: str) -> list:
        raise NotImplementedError


class PdfAdapter(Adapter):
    """PDF = a mix of text and embedded screenshots. Each page -> a text segment;
    each embedded image on the page -> its own image segment, anchored to that page.

    Every failure here RAISES. It must never return a diagnostic string as if it were the
    document's content: downstream, the structuring model treats whatever it is handed as
    source material, so a returned "[PDF parsing needs pypdf]" once became a fully
    invented, published, high-confidence workflow about parsing PDFs. An unreadable source
    has to stop the pipeline, not feed it.
    """
    def decompose(self, data, filename):
        try:
            import pypdf
        except ImportError:
            raise CapabilityUnavailable(
                "PDF text extraction needs pypdf — run: pip install --user pypdf")
        try:
            reader = pypdf.PdfReader(io.BytesIO(data))
        except Exception as e:
            raise CapabilityUnavailable(f"could not open this PDF ({e})")
        if getattr(reader, "is_encrypted", False):
            raise CapabilityUnavailable("this PDF is password-protected")

        out = []
        for i, page in enumerate(reader.pages):
            txt = (page.extract_text() or "").strip()
            if txt:
                out.append({"modality": "pdf_page", "text": txt, "anchor": {"page": i + 1}})
            try:
                for j, img in enumerate(page.images):
                    out.append({"modality": "image", "image_bytes": img.data,
                                "anchor": {"page": i + 1, "image_index": j},
                                "meta": {"from_pdf": True, "name": getattr(img, "name", "")}})
            except Exception:
                pass   # some PDFs have unreadable image streams; text still captured
        if not out:
            raise CapabilityUnavailable(
                "no text or images could be extracted — this looks like a scanned PDF, "
                "which needs OCR")
        return out


class ImageAdapter(Adapter):
    """A single screenshot. A *sequence* of screenshots is just several ImageAdapter
    sources uploaded in order — the pipeline preserves global order across the batch,
    so a 6-screenshot sequence naturally becomes 6 ordered step candidates."""
    def decompose(self, data, filename):
        return [{"modality": "image", "image_bytes": data, "anchor": {}, "meta": {"filename": filename}}]


class TranscriptAdapter(Adapter):
    """Plain text / markdown / a pasted call transcript. Split into coherent chunks;
    keep speaker turns if present."""
    def decompose(self, data, filename):
        text = data.decode("utf-8", "replace") if isinstance(data, bytes) else str(data)
        chunks, buf = [], []
        for line in text.splitlines():
            if line.strip() == "":
                if buf:
                    chunks.append("\n".join(buf).strip())
                    buf = []
            else:
                buf.append(line)
        if buf:
            chunks.append("\n".join(buf).strip())
        return [{"modality": "text", "text": c, "anchor": {"para": i + 1}}
                for i, c in enumerate(chunks) if c] or \
               [{"modality": "text", "text": text.strip(), "anchor": {}}]


class DocxAdapter(Adapter):
    """Word documents. Paragraphs in order, headings kept as headings so the structure
    survives into chunking, and tables flattened to pipe rows — a table read as one blob
    of numbers retrieves badly, a table read as rows retrieves per row."""
    def decompose(self, data, filename):
        try:
            import docx
        except ImportError:
            raise CapabilityUnavailable(
                "DOCX support needs python-docx — run: pip install --user python-docx")
        try:
            doc = docx.Document(io.BytesIO(data))
        except Exception as e:
            raise CapabilityUnavailable(f"could not open this DOCX ({e})")

        out, buf, para_no = [], [], 0
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            style = (p.style.name or "").lower() if p.style is not None else ""
            if style.startswith("heading") and buf:
                para_no += 1
                out.append({"modality": "text", "text": "\n".join(buf),
                            "anchor": {"para": para_no}})
                buf = []
            buf.append(f"## {text}" if style.startswith("heading") else text)
        if buf:
            para_no += 1
            out.append({"modality": "text", "text": "\n".join(buf), "anchor": {"para": para_no}})

        for ti, table in enumerate(doc.tables, 1):
            rows = [" | ".join((c.text or "").strip() for c in r.cells) for r in table.rows]
            rows = [r for r in rows if r.strip(" |")]
            if rows:
                out.append({"modality": "text", "text": "\n".join(rows),
                            "anchor": {"table": ti}, "meta": {"kind": "table"}})

        if not out:
            raise CapabilityUnavailable("this DOCX has no extractable text")
        return out


def _row_segments(header, rows, anchor_key, anchor_value, max_rows=2000):
    """Tabular data → one segment per row, each carrying its column names.

    Row-level granularity is not a detail. Packed as one blob, a model answering "who is
    on call in W32?" has to align columns across lines and reliably blends adjacent rows —
    it returned W33's people for a W32 question. One row per segment makes each row an
    independently retrievable fact and the confusion disappears.

    Wide sheets are batched rather than truncated so a big export degrades in resolution,
    not in coverage.
    """
    header = [h.strip() for h in header]
    body = [r for r in rows if any((c or "").strip() for c in r)]
    per = max(1, -(-len(body) // max_rows))          # rows per segment, ceil division
    out = []
    for i in range(0, len(body), per):
        group = body[i:i + per]
        lines = ["; ".join(f"{h}: {c}" for h, c in zip(header, r)
                           if (c or "").strip() and h)
                 for r in group]
        lines = [ln for ln in lines if ln]
        if not lines:
            continue
        prefix = f"{anchor_key.title()} {anchor_value} — " if anchor_value else ""
        out.append({
            "modality": "text",
            "text": f"{prefix}columns: {', '.join(h for h in header if h)}\n" + "\n".join(lines),
            "anchor": {anchor_key: anchor_value, "row": i + 2},
            "meta": {"kind": "table_row", "rows": len(lines)},
        })
    return out


class SheetAdapter(Adapter):
    """Spreadsheets. One segment per row (see _row_segments), with the header repeated so
    a row retrieved alone still says what its columns mean."""
    MAX_ROWS = 2000

    def decompose(self, data, filename):
        try:
            import openpyxl
        except ImportError:
            raise CapabilityUnavailable(
                "Spreadsheet support needs openpyxl — run: pip install --user openpyxl")
        try:
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        except Exception as e:
            raise CapabilityUnavailable(f"could not open this spreadsheet ({e})")

        out = []
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            header = [str(c) if c is not None else "" for c in rows[0]]
            cells = [[str(c) if c is not None else "" for c in r] for r in rows[1:]]
            out.extend(_row_segments(header, cells, "sheet", ws.title, self.MAX_ROWS))
        wb.close()
        if not out:
            raise CapabilityUnavailable("this spreadsheet has no readable rows")
        return out


class CsvAdapter(Adapter):
    def decompose(self, data, filename):
        import csv as _csv
        text = data.decode("utf-8-sig", "replace") if isinstance(data, bytes) else str(data)
        delim = "\t" if filename.lower().endswith(".tsv") else ","
        rows = list(_csv.reader(io.StringIO(text), delimiter=delim))
        rows = [r for r in rows if any((c or "").strip() for c in r)]
        if not rows:
            raise CapabilityUnavailable("this file has no readable rows")
        out = _row_segments(rows[0], rows[1:], "file", filename)
        if not out:
            raise CapabilityUnavailable("this file has no readable rows")
        return out


class HtmlAdapter(Adapter):
    def decompose(self, data, filename):
        text = data.decode("utf-8", "replace") if isinstance(data, bytes) else str(data)
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(text, "lxml")
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            text = soup.get_text("\n")
        except ImportError:
            text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", text)).strip()
        if not text:
            raise CapabilityUnavailable("no readable text in this page")
        return TranscriptAdapter().decompose(text.encode("utf-8"), filename)


class VideoAdapter(Adapter):
    """Screen recording. Real flow: extract scenes+keyframes (FrameExtractor) and a
    time-stamped transcript (Transcriber), then align them into video_scene segments,
    each carrying its narration text + representative keyframe + [t_start,t_end].
    If those capabilities aren't configured, we raise so the job records 'awaiting
    capability' and keeps the original blob for when a provider is plugged in."""
    def decompose(self, data, filename):
        fx = get_frame_extractor()
        tx = get_transcriber()
        # both raise CapabilityUnavailable until a provider is wired — honest by design
        scenes = fx.scenes("<blob>")           # -> [{t_start,t_end,keyframe_bytes}]
        transcript = tx.transcribe("<blob>", "video/mp4")
        segs = []
        for k, sc in enumerate(scenes):
            narration = _narration_for(transcript, sc["t_start"], sc["t_end"])
            segs.append({"modality": "video_scene", "text": narration,
                         "image_bytes": sc.get("keyframe_bytes"),
                         "anchor": {"t_start": sc["t_start"], "t_end": sc["t_end"]},
                         "meta": {"scene": k}})
        return segs


class AudioAdapter(Adapter):
    def decompose(self, data, filename):
        tx = get_transcriber()
        transcript = tx.transcribe("<blob>", "audio/mp3")   # raises until configured
        return [{"modality": "audio", "text": seg["text"],
                 "anchor": {"t_start": seg["start"], "t_end": seg["end"]}}
                for seg in transcript.get("segments", [])]


def _narration_for(transcript, t0, t1):
    return " ".join(s["text"] for s in transcript.get("segments", [])
                    if s["start"] >= t0 and s["end"] <= t1)


_REGISTRY = {
    "pdf": PdfAdapter(), "docx": DocxAdapter(), "sheet": SheetAdapter(),
    "csv": CsvAdapter(), "html": HtmlAdapter(), "image": ImageAdapter(),
    "transcript": TranscriptAdapter(), "video": VideoAdapter(), "audio": AudioAdapter(),
}


def pick_adapter(mime: str, filename: str) -> Adapter:
    return _REGISTRY[guess_kind(mime, filename)]
